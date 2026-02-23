"""
Data ingestion engine for REXI.
"""

import asyncio
import hashlib
from typing import List, Dict, Optional, Any, BinaryIO
from pathlib import Path
import logging
import mimetypes

from rexi.models.documents import Document, DocumentType
from rexi.models.entities import Entity, EntityType
from rexi.models.relationships import Relationship, RelationshipType
from rexi.services.embedding_service import EmbeddingService
from rexi.services.llm_service import LLMService
from rexi.core.knowledge_graph import KnowledgeGraph
# Temporarily disable entity extractor due to spaCy compatibility issues
# from rexi.agents.entity_extractor import EntityExtractor
# from rexi.agents.relation_extractor import RelationExtractor
from rexi.agents.entity_resolver import EntityResolver

logger = logging.getLogger(__name__)

class IngestionEngine:
    """Engine for ingesting and processing documents."""
    
    def __init__(self):
        """Initialize ingestion engine."""
        self.embedding_service = EmbeddingService()
        self.llm_service = LLMService()
        self.knowledge_graph = KnowledgeGraph()
        # Temporarily disable entity extractor due to spaCy compatibility
        # self.entity_extractor = EntityExtractor()
        # self.relation_extractor = RelationExtractor()
        self.entity_resolver = EntityResolver()
        
        # Supported file types
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.json': self._process_json,
            '.csv': self._process_csv,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.rtf': self._process_text,
            '.epub': self._process_epub,
            '.mobi': self._process_mobi
        }
    
    async def ingest_file(self, file_path: str, source: str = "upload") -> Document:
        """Ingest a single file."""
        try:
            file_path = Path(file_path)
            
            # Check if file exists and is supported
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            extension = file_path.suffix.lower()
            if extension not in self.supported_extensions:
                raise ValueError(f"Unsupported file type: {extension}")
            
            # Create document record
            document = Document(
                title=file_path.stem,
                type=self._get_document_type(extension),
                content="",  # Will be filled by processor
                file_path=str(file_path),
                file_size=file_path.stat().st_size,
                mime_type=mimetypes.guess_type(str(file_path))[0],
                source=source
            )
            
            # Process file content
            processor = self.supported_extensions[extension]
            content = await processor(file_path)
            document.content = content
            document.word_count = len(content.split())
            
            # Detect language
            try:
                import langdetect
                document.language = langdetect.detect(content)
            except:
                document.language = "en"
            
            # Generate embedding for the document
            if content:
                document.embedding = self.embedding_service.encode_text(content)
            
            # Extract entities and relationships
            if self.llm_service.is_available():
                await self._extract_knowledge(document)
            
            document.mark_processed()
            logger.info(f"Successfully ingested: {file_path}")
            
            return document
            
        except Exception as e:
            logger.error(f"Failed to ingest file {file_path}: {e}")
            raise
    
    async def ingest_text(self, text: str, title: str = "", source: str = "text_input") -> Document:
        """Ingest text directly."""
        try:
            document = Document(
                title=title or f"Text_{hashlib.md5(text.encode()).hexdigest()[:8]}",
                type=DocumentType.TEXT,
                content=text,
                source=source,
                word_count=len(text.split())
            )
            
            # Detect language
            try:
                import langdetect
                document.language = langdetect.detect(text)
            except:
                document.language = "en"
            
            # Generate embedding
            document.embedding = self.embedding_service.encode_text(text)
            
            # Extract entities and relationships
            if self.llm_service.is_available():
                await self._extract_knowledge(document)
            
            document.mark_processed()
            logger.info(f"Successfully ingested text: {title}")
            
            return document
            
        except Exception as e:
            logger.error(f"Failed to ingest text: {e}")
            raise
    
    async def _extract_knowledge(self, document: Document):
        """Extract entities and relationships from document."""
        try:
            # Split content into chunks for processing
            chunk_size = 4000  # Tokens
            chunks = self._chunk_text(document.content, chunk_size)
            
            all_entities = []
            all_relationships = []
            
            for i, chunk in enumerate(chunks):
                # Temporarily disable entity extraction due to spaCy compatibility
                # entities_data = self.entity_extractor.extract_entities(chunk)
                entities_data = []
                
                # Temporarily disable relation extraction due to spaCy compatibility
                # relationships_data = self.relation_extractor.extract_relations(chunk, entities_data)
                relationships_data = []
                
                # Process entities
                for entity_data in entities_data:
                    entity = Entity(
                        name=entity_data["text"],
                        type=entity_data["type"],
                        description=entity_data.get("context", ""),
                        confidence=entity_data["confidence"],
                        source_references=[document.id],
                        properties={
                            "chunk_index": i,
                            "extraction_method": entity_data["source"],
                            "context": entity_data.get("context", "")
                        }
                    )
                    
                    # Generate embedding for entity
                    entity_text = f"{entity.name} {entity.description or ''}"
                    entity.embedding = self.embedding_service.encode_text(entity_text)
                    
                    all_entities.append(entity)
                
                # Process relationships
                for rel_data in relationships_data:
                    relationship = Relationship(
                        source_entity_id=rel_data["source"],
                        target_entity_id=rel_data["target"],
                        type=rel_data["type"],
                        confidence=rel_data["confidence"],
                        evidence_references=[document.id],
                        properties={
                            "extraction_method": rel_data["source_method"],
                            "evidence": rel_data.get("evidence", ""),
                            "chunk_index": i
                        }
                    )
                    all_relationships.append(relationship)
            
            # Resolve entities (deduplication and merging)
            resolved_entities = self.entity_resolver.resolve_entities(all_entities)
            
            # Add entities to knowledge graph
            entity_id_mapping = {}
            for entity in resolved_entities:
                stored_id = self.knowledge_graph.add_entity(entity)
                entity_id_mapping[entity.name] = stored_id
            
            # Update relationships with resolved entity IDs
            resolved_relationships = []
            for relationship in all_relationships:
                # Map entity names to IDs
                source_id = entity_id_mapping.get(relationship.source_entity_id)
                target_id = entity_id_mapping.get(relationship.target_entity_id)
                
                if source_id and target_id:
                    relationship.source_entity_id = source_id
                    relationship.target_entity_id = target_id
                    resolved_relationships.append(relationship)
            
            # Add relationships to knowledge graph
            for relationship in resolved_relationships:
                self.knowledge_graph.add_relationship(relationship)
            
            logger.info(f"Extracted {len(resolved_entities)} entities and {len(resolved_relationships)} relationships")
            
        except Exception as e:
            logger.error(f"Knowledge extraction failed: {e}")
            document.mark_error(str(e))
    
    def _chunk_text(self, text: str, max_tokens: int = 4000) -> List[str]:
        """Split text into chunks of maximum token size."""
        if not self.llm_service.tokenizer:
            return [text]
        
        chunks = []
        current_chunk = ""
        current_tokens = 0
        
        sentences = text.split('. ')
        
        for sentence in sentences:
            sentence_tokens = self.llm_service.count_tokens(sentence)
            
            if current_tokens + sentence_tokens > max_tokens and current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = sentence
                current_tokens = sentence_tokens
            else:
                current_chunk += sentence + ". "
                current_tokens += sentence_tokens
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _normalize_entity_name(self, name: str) -> str:
        """Normalize entity name for deduplication."""
        return name.lower().strip().replace(" ", "_")
    
    def _get_document_type(self, extension: str) -> DocumentType:
        """Map file extension to document type."""
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.txt': DocumentType.TEXT,
            '.md': DocumentType.MARKDOWN,
            '.json': DocumentType.TEXT,
            '.csv': DocumentType.TEXT,
            '.docx': DocumentType.TEXT,
            '.doc': DocumentType.TEXT,
            '.rtf': DocumentType.TEXT,
            '.epub': DocumentType.TEXT,
            '.mobi': DocumentType.TEXT
        }
        return type_mapping.get(extension, DocumentType.TEXT)
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF file."""
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(str(file_path))
            text = ""
            
            for page in doc:
                text += page.get_text()
            
            doc.close()
            return text
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            return ""
    
    async def _process_text(self, file_path: Path) -> str:
        """Process plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Text processing failed: {e}")
            return ""
    
    async def _process_markdown(self, file_path: Path) -> str:
        """Process Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Markdown processing failed: {e}")
            return ""
    
    async def _process_json(self, file_path: Path) -> str:
        """Process JSON file."""
        try:
            import json
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return json.dumps(data, indent=2)
        except Exception as e:
            logger.error(f"JSON processing failed: {e}")
            return ""
    
    async def _process_csv(self, file_path: Path) -> str:
        """Process CSV file."""
        try:
            import pandas as pd
            df = pd.read_csv(file_path)
            return df.to_string()
        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            return ""
    
    async def _process_docx(self, file_path: Path) -> str:
        """Process DOCX file."""
        try:
            from docx import Document
            doc = Document(str(file_path))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            return ""
    
    async def _process_epub(self, file_path: Path) -> str:
        """Process EPUB file."""
        try:
            import ebooklib
            from ebooklib import epub
            book = epub.read_epub(str(file_path))
            text = ""
            
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    text += item.get_content().decode('utf-8') + "\n"
            
            return text
        except Exception as e:
            logger.error(f"EPUB processing failed: {e}")
            return ""
    
    async def _process_mobi(self, file_path: Path) -> str:
        """Process MOBI file."""
        try:
            # MOBI processing requires additional libraries
            # For now, return empty string
            logger.warning("MOBI processing not implemented yet")
            return ""
        except Exception as e:
            logger.error(f"MOBI processing failed: {e}")
            return ""
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions."""
        return list(self.supported_extensions.keys())
    
    def close(self):
        """Close ingestion engine."""
        try:
            self.knowledge_graph.close()
            logger.info("Ingestion engine closed")
        except Exception as e:
            logger.error(f"Error closing ingestion engine: {e}")
